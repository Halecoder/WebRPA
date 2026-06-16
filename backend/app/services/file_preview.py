"""文件预览服务 - 支持 Excel、Word、PPT、PDF 等文档的原生预览"""
import os
import io
import html
import hashlib
import tempfile
from pathlib import Path
from typing import Optional, Tuple
import base64


# 预览缓存目录
_preview_cache_dir: Optional[Path] = None


def get_preview_cache_dir() -> Path:
    """获取预览缓存目录"""
    global _preview_cache_dir
    if _preview_cache_dir is None:
        _preview_cache_dir = Path(tempfile.gettempdir()) / "webrpa_file_preview"
        _preview_cache_dir.mkdir(exist_ok=True)
    return _preview_cache_dir


def get_cache_key(file_path: Path) -> str:
    """生成文件缓存键"""
    file_hash = hashlib.md5(str(file_path).encode()).hexdigest()
    mtime = int(file_path.stat().st_mtime)
    return f"{file_hash}_{mtime}"


def preview_excel(file_path: Path) -> Tuple[str, str]:
    """
    预览 Excel 文件，返回 (html_content, content_type)
    支持 .xlsx, .xls, .csv
    """
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.csv':
            return _preview_csv(file_path)
        elif ext == '.xlsx':
            return _preview_xlsx(file_path)
        elif ext == '.xls':
            return _preview_xls(file_path)
        else:
            return f"<p>不支持的 Excel 格式: {ext}</p>", "text/html"
    except Exception as e:
        return f"<p>预览失败: {html.escape(str(e))}</p>", "text/html"


def _preview_csv(file_path: Path) -> Tuple[str, str]:
    """预览 CSV 文件"""
    import csv
    
    rows = []
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)[:1000]  # 限制行数
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    if not rows:
        return "<p>无法读取 CSV 文件</p>", "text/html"
    
    return _generate_table_html(rows, file_path.name), "text/html"


def _preview_xlsx(file_path: Path) -> Tuple[str, str]:
    """预览 XLSX 文件"""
    import openpyxl
    
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sheets_html = []
    
    for sheet_name in wb.sheetnames[:10]:  # 限制工作表数量
        ws = wb[sheet_name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 1000:  # 限制行数
                break
            rows.append(list(row))
        
        if rows:
            sheet_html = f'<div class="sheet"><h3>{html.escape(sheet_name)}</h3>'
            sheet_html += _generate_table_html(rows, None, show_title=False)
            sheet_html += '</div>'
            sheets_html.append(sheet_html)
    
    wb.close()
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="sheets">' + ''.join(sheets_html) + '</div>',
        extra_style='.sheet { margin-bottom: 24px; } .sheet h3 { margin-bottom: 12px; color: #3b82f6; }'
    )
    return content, "text/html"


def _preview_xls(file_path: Path) -> Tuple[str, str]:
    """预览 XLS 文件"""
    import xlrd
    
    wb = xlrd.open_workbook(file_path)
    sheets_html = []
    
    for sheet_idx in range(min(wb.nsheets, 10)):
        ws = wb.sheet_by_index(sheet_idx)
        rows = []
        for row_idx in range(min(ws.nrows, 1000)):
            row = [ws.cell_value(row_idx, col_idx) for col_idx in range(ws.ncols)]
            rows.append(row)
        
        if rows:
            sheet_html = f'<div class="sheet"><h3>{html.escape(ws.name)}</h3>'
            sheet_html += _generate_table_html(rows, None, show_title=False)
            sheet_html += '</div>'
            sheets_html.append(sheet_html)
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="sheets">' + ''.join(sheets_html) + '</div>',
        extra_style='.sheet { margin-bottom: 24px; } .sheet h3 { margin-bottom: 12px; color: #3b82f6; }'
    )
    return content, "text/html"


def _generate_table_html(rows: list, filename: Optional[str], show_title: bool = True) -> str:
    """生成表格 HTML"""
    if not rows:
        return "<p>空表格</p>"
    
    table_html = '<div class="table-wrapper"><table>'
    
    # 表头
    if rows:
        table_html += '<thead><tr>'
        for cell in rows[0]:
            cell_str = str(cell) if cell is not None else ''
            table_html += f'<th>{html.escape(cell_str)}</th>'
        table_html += '</tr></thead>'
    
    # 表体
    table_html += '<tbody>'
    for row in rows[1:]:
        table_html += '<tr>'
        for cell in row:
            cell_str = str(cell) if cell is not None else ''
            table_html += f'<td>{html.escape(cell_str)}</td>'
        table_html += '</tr>'
    table_html += '</tbody></table></div>'
    
    if show_title and filename:
        return _get_preview_wrapper(filename, table_html)
    return table_html


def _convert_legacy_office(file_path: Path, target_ext: str) -> Path | None:
    """用 win32com 把旧版 .doc/.ppt/.xls 转为新格式（需安装对应 Office）。失败返回 None。"""
    import tempfile
    ext = file_path.suffix.lower()
    out = Path(tempfile.gettempdir()) / (file_path.stem + "_conv" + target_ext)
    try:
        import win32com.client as win32  # type: ignore
        if ext == '.doc':
            app = win32.Dispatch('Word.Application')
            app.Visible = False
            try:
                doc = app.Documents.Open(str(file_path))
                doc.SaveAs(str(out), FileFormat=16)  # wdFormatDocumentDefault(.docx)
                doc.Close()
            finally:
                app.Quit()
            return out if out.exists() else None
        if ext == '.ppt':
            app = win32.Dispatch('PowerPoint.Application')
            try:
                pres = app.Presentations.Open(str(file_path), WithWindow=False)
                pres.SaveAs(str(out), 24)  # ppSaveAsOpenXMLPresentation(.pptx)
                pres.Close()
            finally:
                app.Quit()
            return out if out.exists() else None
    except Exception as e:
        print(f"[file_preview] 旧版 Office 转换失败: {e}")
    return None


def preview_word(file_path: Path) -> Tuple[str, str]:
    """
    预览 Word 文件，返回 (html_content, content_type)
    支持 .docx；.doc 自动尝试转换为 .docx（需安装 Word）
    """
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.docx':
            return _preview_docx(file_path)
        elif ext == '.doc':
            converted = _convert_legacy_office(file_path, '.docx')
            if converted:
                try:
                    return _preview_docx(converted)
                finally:
                    try:
                        converted.unlink()
                    except Exception:
                        pass
            return "<p>无法预览 .doc：未检测到可用的 Word 组件。请转换为 .docx 后重试。</p>", "text/html"
        else:
            return f"<p>不支持的 Word 格式: {ext}</p>", "text/html"
    except Exception as e:
        return f"<p>预览失败: {html.escape(str(e))}</p>", "text/html"


def _preview_docx(file_path: Path) -> Tuple[str, str]:
    """预览 DOCX 文件"""
    from docx import Document
    from docx.shared import Inches
    
    doc = Document(file_path)
    content_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            # 检查段落样式
            style_name = para.style.name if para.style else ''
            if 'Heading' in style_name or 'heading' in style_name.lower() or '标题' in style_name:
                level = 2
                if '1' in style_name:
                    level = 1
                elif '2' in style_name:
                    level = 2
                elif '3' in style_name:
                    level = 3
                content_parts.append(f'<h{level}>{html.escape(para.text)}</h{level}>')
            elif 'Title' in style_name or 'title' in style_name.lower():
                content_parts.append(f'<h1 class="doc-title">{html.escape(para.text)}</h1>')
            else:
                content_parts.append(f'<p>{html.escape(para.text)}</p>')
    
    # 处理表格
    for table in doc.tables:
        table_html = '<table class="doc-table">'
        for i, row in enumerate(table.rows):
            table_html += '<tr>'
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                table_html += f'<{tag}>{html.escape(cell.text)}</{tag}>'
            table_html += '</tr>'
        table_html += '</table>'
        content_parts.append(table_html)
    
    if not content_parts:
        content_parts.append('<p class="empty-doc">文档内容为空</p>')
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="doc-content">' + ''.join(content_parts) + '</div>',
        extra_style='''
            .doc-content { max-width: 800px; margin: 0 auto; line-height: 1.8; }
            .doc-content p { margin-bottom: 12px; text-indent: 2em; }
            .doc-content h1, .doc-content h2, .doc-content h3 { margin: 20px 0 12px; text-indent: 0; }
            .doc-content .doc-title { text-align: center; font-size: 24px; margin-bottom: 24px; }
            .doc-table { width: 100%; margin: 16px 0; border-collapse: collapse; }
            .doc-table th, .doc-table td { padding: 10px 12px; border: 1px solid #3f3f46; text-indent: 0; }
            .doc-table th { background: #27272a; font-weight: 600; }
            .empty-doc { text-align: center; color: #71717a; }
        '''
    )
    return content, "text/html"


def preview_ppt(file_path: Path) -> Tuple[str, str]:
    """
    预览 PPT 文件，返回 (html_content, content_type)
    支持 .pptx
    """
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.pptx':
            return _preview_pptx(file_path)
        elif ext == '.ppt':
            converted = _convert_legacy_office(file_path, '.pptx')
            if converted:
                try:
                    return _preview_pptx(converted)
                finally:
                    try:
                        converted.unlink()
                    except Exception:
                        pass
            return "<p>无法预览 .ppt：未检测到可用的 PowerPoint 组件。请转换为 .pptx 后重试。</p>", "text/html"
        else:
            return f"<p>不支持的 PPT 格式: {ext}</p>", "text/html"
    except ImportError:
        return "<p>预览 PPT 需要安装 python-pptx 库: pip install python-pptx</p>", "text/html"
    except Exception as e:
        return f"<p>预览失败: {html.escape(str(e))}</p>", "text/html"


def _preview_pptx(file_path: Path) -> Tuple[str, str]:
    """预览 PPTX 文件 - 使用 zipfile 直接解析 XML"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    slides_html = []
    
    # PPTX 是一个 ZIP 文件，直接解析 XML
    with zipfile.ZipFile(file_path, 'r') as zf:
        # 获取所有幻灯片文件
        slide_files = sorted([f for f in zf.namelist() if f.startswith('ppt/slides/slide') and f.endswith('.xml')])
        
        for idx, slide_file in enumerate(slide_files[:50]):
            slide_content = []
            slide_content.append(f'<div class="slide-number">第 {idx + 1} 页</div>')
            
            try:
                with zf.open(slide_file) as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    
                    # 提取所有文本
                    texts = []
                    for elem in root.iter():
                        # 查找文本元素 (a:t)
                        if elem.tag.endswith('}t') and elem.text:
                            texts.append(elem.text)
                    
                    # 合并连续文本，按段落分组
                    if texts:
                        # 第一个通常是标题
                        if len(texts) > 0:
                            slide_content.append(f'<h2>{html.escape(texts[0])}</h2>')
                        # 其余是内容
                        for text in texts[1:]:
                            if text.strip():
                                slide_content.append(f'<p>{html.escape(text)}</p>')
            except Exception:
                pass
            
            if len(slide_content) == 1:
                slide_content.append('<p class="empty-slide">（此页无文本内容）</p>')
            
            slides_html.append(f'<div class="slide">{"".join(slide_content)}</div>')
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="slides">' + ''.join(slides_html) + '</div>',
        extra_style='''
            .slides { display: flex; flex-direction: column; gap: 24px; }
            .slide { background: #27272a; border-radius: 8px; padding: 24px; border: 1px solid #3f3f46; }
            .slide-number { font-size: 12px; color: #71717a; margin-bottom: 12px; }
            .slide h2 { font-size: 20px; margin-bottom: 16px; color: #3b82f6; }
            .slide p { margin-bottom: 8px; line-height: 1.6; }
            .empty-slide { color: #71717a; font-style: italic; }
        '''
    )
    return content, "text/html"


def preview_pdf_as_images(file_path: Path) -> Tuple[str, str]:
    """
    将 PDF 转换为图片预览，返回 (html_content, content_type)
    使用 pdf2image 库替代 PyMuPDF
    """
    try:
        from pdf2image import convert_from_path
        
        # 获取 backend 目录的路径
        backend_root = Path(__file__).parent.parent.parent
        poppler_path = backend_root / 'poppler' / 'Library' / 'bin'
        
        # 转换 PDF 为图片，限制页数
        if poppler_path.exists():
            images = convert_from_path(str(file_path), dpi=150, first_page=1, last_page=50, poppler_path=str(poppler_path))
        else:
            # 如果内置 poppler 不存在，尝试使用系统 PATH 中的 poppler
            images = convert_from_path(str(file_path), dpi=150, first_page=1, last_page=50)
        
        pages_html = []
        
        for page_num, img in enumerate(images):
            # 将图片转换为 base64
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_bytes = img_buffer.getvalue()
            b64_image = base64.b64encode(img_bytes).decode('utf-8')
            
            pages_html.append(f'''
                <div class="pdf-page">
                    <div class="page-number">第 {page_num + 1} 页</div>
                    <img src="data:image/png;base64,{b64_image}">
                </div>
            ''')
        
        content = _get_preview_wrapper(
            file_path.name,
            '<div class="pdf-pages">' + ''.join(pages_html) + '</div>',
            extra_style='''
                .pdf-pages { display: flex; flex-direction: column; align-items: center; gap: 16px; }
                .pdf-page { background: #27272a; border-radius: 8px; padding: 16px; text-align: center; }
                .page-number { font-size: 12px; color: #71717a; margin-bottom: 8px; }
                .pdf-page img { max-width: 100%; border-radius: 4px; box-shadow: 0 4px 12px rgba(0,0,0,0.3); }
            '''
        )
        return content, "text/html"
    except ImportError:
        error_html = "<p>PDF 预览需要安装 pdf2image 库: pip install pdf2image</p>"
        return error_html, "text/html"
    except Exception as e:
        error_msg = str(e)
        if "poppler" in error_msg.lower():
            error_html = "<div style='padding: 20px; color: #fafafa;'>"
            error_html += "<h2 style='color: #ef4444;'>PDF 预览需要 poppler 工具</h2>"
            error_html += "<p>请将 poppler 文件夹放置到 backend 目录下</p>"
            error_html += "<p>下载地址: <a href='https://github.com/oschwartz10612/poppler-windows/releases' target='_blank' style='color: #3b82f6;'>https://github.com/oschwartz10612/poppler-windows/releases</a></p>"
            error_html += f"<p style='color: #71717a; margin-top: 20px;'>错误详情: {html.escape(error_msg)}</p>"
            error_html += "</div>"
            return error_html, "text/html"
        return f"<p>PDF 预览失败: {html.escape(error_msg)}</p>", "text/html"


def _get_preview_wrapper(filename: str, content: str, extra_style: str = '') -> str:
    """生成预览页面包装器"""
    return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(filename)} - 预览</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            background: #18181b; 
            color: #fafafa; 
            padding: 20px;
            min-height: 100vh;
        }}
        .header {{
            text-align: center;
            margin-bottom: 24px;
            padding-bottom: 16px;
            border-bottom: 1px solid #27272a;
        }}
        .header h1 {{
            font-size: 18px;
            font-weight: 600;
            color: #fafafa;
            word-break: break-all;
        }}
        .table-wrapper {{
            overflow-x: auto;
            margin: 16px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
        }}
        th, td {{
            padding: 10px 12px;
            text-align: left;
            border: 1px solid #3f3f46;
            white-space: nowrap;
            max-width: 300px;
            overflow: hidden;
            text-overflow: ellipsis;
        }}
        th {{
            background: #27272a;
            font-weight: 600;
            position: sticky;
            top: 0;
        }}
        tr:hover td {{
            background: #27272a;
        }}
        {extra_style}
    </style>
</head>
<body>
    <div class="header">
        <h1>📄 {html.escape(filename)}</h1>
    </div>
    {content}
</body>
</html>'''


def get_preview_content(file_path: Path) -> Optional[Tuple[bytes, str]]:
    """
    获取文件预览内容
    返回 (content_bytes, content_type) 或 None
    """
    ext = file_path.suffix.lower()
    
    # Excel 文件
    if ext in ['.xlsx', '.xls', '.csv']:
        content, content_type = preview_excel(file_path)
        return content.encode('utf-8'), content_type
    
    # Word 文件
    if ext in ['.docx', '.doc']:
        content, content_type = preview_word(file_path)
        return content.encode('utf-8'), content_type
    
    # PPT 文件
    if ext in ['.pptx', '.ppt']:
        content, content_type = preview_ppt(file_path)
        return content.encode('utf-8'), content_type
    
    return None


def is_previewable_document(filename: str) -> bool:
    """检查文件是否支持文档预览"""
    ext = Path(filename).suffix.lower()
    return ext in ['.xlsx', '.xls', '.csv', '.docx', '.doc', '.pptx', '.ppt']
