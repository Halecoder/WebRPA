"""WebRPA小助手 - HTTP API"""
from __future__ import annotations

from typing import Any

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel

from app.models.ai_assistant import (
    AssistantConfig,
    ChatRequest,
    ChatResponse,
    CreateSessionRequest,
    SessionListItem,
)
from app.services.ai_assistant_service import (
    chat_once,
    create_session,
    delete_session,
    list_sessions,
    load_session,
    save_session,
    cancel_session,
)
from app.services.ai_assistant_skills import registry as skill_registry


router = APIRouter(prefix="/api/ai-assistant", tags=["ai-assistant"])


# Socket.IO 实例（运行时由 main.py 注入）
sio = None


def set_sio(socketio_instance) -> None:
    global sio
    sio = socketio_instance


# ---------- 会话管理 ----------

@router.get("/sessions")
async def api_list_sessions() -> list[SessionListItem]:
    return list_sessions()


@router.post("/sessions")
async def api_create_session(req: CreateSessionRequest):
    session = create_session(req.title)
    return {"session_id": session.id, "title": session.title}


@router.get("/sessions/{session_id}")
async def api_get_session(session_id: str):
    session = load_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail=f"会话不存在: {session_id}")
    return session


@router.delete("/sessions/{session_id}")
async def api_delete_session(session_id: str):
    ok = delete_session(session_id)
    if not ok:
        raise HTTPException(status_code=404, detail=f"会话不存在: {session_id}")
    return {"success": True}


class RenameSessionRequest(BaseModel):
    title: str


@router.patch("/sessions/{session_id}/title")
async def api_rename_session(session_id: str, req: RenameSessionRequest):
    session = load_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail=f"会话不存在: {session_id}")
    session.title = (req.title or "").strip() or session.title
    save_session(session)
    return {"success": True, "title": session.title}


# ---------- 对话 ----------

@router.post("/chat", response_model=ChatResponse)
async def api_chat(req: ChatRequest):
    """一次性对话（非流式）"""
    # 加载或创建会话
    session = None
    if req.session_id:
        session = load_session(req.session_id)
    if session is None:
        # honor 前端预先生成的 session_id（便于执行期间精准取消）
        session = create_session(session_id=req.session_id)

    sid_for_emit = session.id

    async def on_event(event_name: str, payload: dict[str, Any]):
        if sio is None:
            return
        try:
            await sio.emit(f"ai_assistant:{event_name}", {"session_id": sid_for_emit, **payload})
        except Exception as e:
            print(f"[AIAssistant] 推送事件失败 {event_name}: {e}")

    session = await chat_once(
        session=session,
        user_message_text=req.message,
        config=req.config,
        workflow_context=req.workflow_context,
        images=req.images,
        on_event=on_event,
    )

    # 返回最后一条助手消息
    last_assistant = None
    for m in reversed(session.messages):
        if m.role.value == "assistant" and not m.tool_calls:
            last_assistant = m
            break
    if last_assistant is None:
        # 兜底：返回最后一条
        last_assistant = session.messages[-1] if session.messages else None
    if last_assistant is None:
        raise HTTPException(status_code=500, detail="助手未返回任何消息")

    return ChatResponse(session_id=session.id, message=last_assistant)


@router.post("/sessions/{session_id}/cancel")
async def api_cancel_session(session_id: str):
    """中断当前会话正在跑的对话/工具任务"""
    ok = cancel_session(session_id)
    return {"success": ok, "session_id": session_id}


# ---------- 连通性测试 ----------

class TestConnectionRequest(BaseModel):
    config: AssistantConfig


@router.post("/test-connection")
async def api_test_connection(req: TestConnectionRequest):
    """用极简请求测试 模型/密钥/地址 是否正确可用。

    返回 {success, message, detail, latency_ms}。不写会话、不带工具，
    只发一条最短消息，快速验证配置（能当场发现填错的地址/模型/密钥/多余空格等）。
    """
    import time as _time
    from app.services.ai_assistant_service import _call_llm, LLMError, _normalize_api_url

    cfg = req.config
    # 基础校验
    if not (cfg.api_url or "").strip():
        return {"success": False, "message": "API 地址不能为空", "detail": ""}
    if not (cfg.model or "").strip():
        return {"success": False, "message": "模型名称不能为空", "detail": ""}

    try:
        resolved_url = _normalize_api_url(cfg.api_url)
    except Exception as e:
        return {"success": False, "message": f"API 地址无效：{e}", "detail": ""}

    t0 = _time.time()
    try:
        # 非流式（on_event=None）、不带工具、最短消息，快速探活
        data = await _call_llm(
            config=cfg,
            messages=[{"role": "user", "content": "ping"}],
            tools=None,
            on_event=None,
        )
        latency = int((_time.time() - t0) * 1000)
        # 解析返回的模型名/内容做轻量校验
        model_echo = ""
        try:
            model_echo = data.get("model") or ""
        except Exception:
            pass
        return {
            "success": True,
            "message": "连接成功，配置可用",
            "detail": f"实际请求地址：{resolved_url}" + (f"；返回模型：{model_echo}" if model_echo else ""),
            "latency_ms": latency,
        }
    except LLMError as e:
        return {
            "success": False,
            "message": str(e)[:500],
            "detail": f"实际请求地址：{resolved_url}",
            "latency_ms": int((_time.time() - t0) * 1000),
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"测试失败：{type(e).__name__}: {str(e)[:400]}",
            "detail": f"实际请求地址：{resolved_url}",
        }


# ---------- 附件文本提取（多模态/文档理解） ----------

class ExtractFileRequest(BaseModel):
    filename: str
    content_base64: str  # 可为 data URL 或纯 base64


def _decode_to_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


@router.post("/extract-file")
async def api_extract_file(req: ExtractFileRequest):
    """把用户上传的文档（pdf/docx/xls/xlsx/csv/txt/md/html 等）提取成纯文本，
    供小助手（多模态/文档理解）分析。图片不走这里（走 chat 的 images 字段）。
    """
    import base64 as _b64
    import os as _os
    import tempfile as _tmp

    name = (req.filename or "file").strip()
    ext = _os.path.splitext(name)[1].lower()
    raw = req.content_base64 or ""
    if "," in raw and raw.strip().startswith("data:"):
        raw = raw.split(",", 1)[1]
    try:
        data = _b64.b64decode(raw)
    except Exception as e:
        return {"success": False, "error": f"文件解码失败: {e}", "text": ""}

    MAX_CHARS = 30000

    def _clip(t: str) -> str:
        t = t or ""
        if len(t) > MAX_CHARS:
            return t[:MAX_CHARS] + f"\n…（内容过长，已截断，仅保留前 {MAX_CHARS} 字）"
        return t

    try:
        # 纯文本类：直接解码
        if ext in (".txt", ".md", ".markdown", ".csv", ".html", ".htm", ".svg", ".json", ".log", ".xml", ".yaml", ".yml", ".ini", ".tsv"):
            return {"success": True, "text": _clip(_decode_to_text(data)), "error": ""}

        # 写临时文件供二进制解析库使用
        fd, tmp_path = _tmp.mkstemp(suffix=ext or ".bin")
        _os.close(fd)
        try:
            with open(tmp_path, "wb") as f:
                f.write(data)

            if ext == ".pdf":
                try:
                    import pdfplumber
                    parts: list[str] = []
                    with pdfplumber.open(tmp_path) as pdf:
                        for i, page in enumerate(pdf.pages):
                            if i >= 50:
                                parts.append("…（页数过多，仅解析前 50 页）")
                                break
                            parts.append(page.extract_text() or "")
                    return {"success": True, "text": _clip("\n".join(parts)), "error": ""}
                except Exception:
                    from pypdf import PdfReader
                    reader = PdfReader(tmp_path)
                    parts = [(p.extract_text() or "") for p in reader.pages[:50]]
                    return {"success": True, "text": _clip("\n".join(parts)), "error": ""}

            if ext == ".docx":
                from docx import Document
                doc = Document(tmp_path)
                parts = [p.text for p in doc.paragraphs]
                # 表格内容
                for tbl in doc.tables:
                    for row in tbl.rows:
                        parts.append("\t".join(c.text for c in row.cells))
                return {"success": True, "text": _clip("\n".join(parts)), "error": ""}

            if ext == ".xlsx":
                import openpyxl
                wb = openpyxl.load_workbook(tmp_path, read_only=True, data_only=True)
                out: list[str] = []
                for ws in wb.worksheets:
                    out.append(f"# 工作表: {ws.title}")
                    for ri, row in enumerate(ws.iter_rows(values_only=True)):
                        if ri >= 500:
                            out.append("…（行数过多，仅解析前 500 行）")
                            break
                        out.append("\t".join("" if c is None else str(c) for c in row))
                wb.close()
                return {"success": True, "text": _clip("\n".join(out)), "error": ""}

            if ext == ".xls":
                import xlrd
                wb = xlrd.open_workbook(tmp_path)
                out = []
                for ws in wb.sheets():
                    out.append(f"# 工作表: {ws.name}")
                    for ri in range(min(ws.nrows, 500)):
                        out.append("\t".join(str(v) for v in ws.row_values(ri)))
                return {"success": True, "text": _clip("\n".join(out)), "error": ""}

            if ext == ".doc":
                return {"success": False, "error": "暂不支持旧版 .doc（请另存为 .docx 后再上传）", "text": ""}

            # 兜底：尝试当文本解码
            return {"success": True, "text": _clip(_decode_to_text(data)), "error": ""}
        finally:
            try:
                _os.remove(tmp_path)
            except Exception:
                pass
    except Exception as e:
        return {"success": False, "error": f"提取失败: {type(e).__name__}: {str(e)[:200]}", "text": ""}


# ---------- Skills 元数据 ----------

@router.get("/skills")
async def api_list_skills():
    """列出所有可用的 Skills（前端用于渲染调试面板）"""
    return {
        "count": len(skill_registry.names()),
        "skills": [
            {
                "name": s.name,
                "description": s.description,
                "parameters": s.parameters,
                "requires_approval": s.requires_approval,
            }
            for s in skill_registry._skills.values()  # noqa: SLF001
        ],
    }


# ---------- 长期记忆 ----------

@router.get("/memories")
async def api_list_memories():
    from app.services.ai_assistant_skills import _load_memory  # type: ignore
    return {"entries": _load_memory()}


class AddMemoryRequest(BaseModel):
    content: str
    tags: list[str] = []


@router.post("/memories")
async def api_add_memory(req: AddMemoryRequest):
    from app.services.ai_assistant_skills import skill_remember
    return await skill_remember(content=req.content, tags=req.tags)


@router.delete("/memories/{memory_id}")
async def api_delete_memory(memory_id: str):
    from app.services.ai_assistant_skills import skill_forget
    return await skill_forget(memory_id=memory_id)


# ============================================================
# MCP（Model Context Protocol）相关 API
# ============================================================

class MCPConfigPayload(BaseModel):
    config: dict  # {"mcpServers": {...}}


@router.get("/mcp/config")
async def api_get_mcp_config():
    """读取当前 MCP 配置（mcpServers JSON）"""
    from app.services.mcp_manager import load_mcp_config
    return load_mcp_config()


@router.put("/mcp/config")
async def api_save_mcp_config(req: MCPConfigPayload):
    """覆盖保存 MCP 配置（不会自动重连，需要再调 /mcp/reload）"""
    from app.services.mcp_manager import save_mcp_config
    if not isinstance(req.config, dict):
        raise HTTPException(400, "config 必须是 JSON 对象")
    if "mcpServers" not in req.config:
        # 兼容用户直接传 {"server-name": {...}, ...}
        req.config = {"mcpServers": req.config}
    save_mcp_config(req.config)
    return {"success": True, "saved": True}


@router.post("/mcp/reload")
async def api_reload_mcp():
    """根据当前配置重新连接所有 MCP 服务器（保存配置后调一次让改动生效）"""
    from app.services.mcp_manager import mcp_manager
    return await mcp_manager.reload()


@router.get("/mcp/status")
async def api_mcp_status():
    """查看所有 MCP 服务器的连接状态、工具列表"""
    from app.services.mcp_manager import mcp_manager
    return mcp_manager.status()


@router.post("/mcp/disconnect-all")
async def api_mcp_disconnect_all():
    """断开所有 MCP 服务器（用于排错或临时禁用）"""
    from app.services.mcp_manager import mcp_manager
    await mcp_manager.shutdown_all()
    return {"success": True, "disconnected": True}
